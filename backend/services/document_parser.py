import io
import zipfile
from functools import lru_cache
from pathlib import Path
from typing import Any

from docx import Document
from PyPDF2 import PdfReader


MAX_ARCHIVE_ENTRIES = 4_096
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_ARCHIVE_COMPRESSION_RATIO = 200
MAX_PDF_PAGES = 300
MAX_OCR_PIXELS = 20_000_000
MAX_XLSX_ROWS = 100_000


def validate_document_signature(filename: str, content: bytes) -> None:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" and not content.startswith(b"%PDF-"):
        raise ValueError("文件扩展名为 PDF，但内容签名不匹配。")
    if suffix in {".docx", ".xlsx"}:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                entries = archive.infolist()
        except zipfile.BadZipFile as exc:
            raise ValueError(f"文件扩展名为 {suffix[1:].upper()}，但不是有效的 Office 文档。") from exc
        if len(entries) > MAX_ARCHIVE_ENTRIES:
            raise ValueError("Office 文档包含过多压缩条目。")
        total_size = sum(entry.file_size for entry in entries)
        if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
            raise ValueError("Office 文档解压后过大。")
        if any(entry.compress_size and entry.file_size / entry.compress_size > MAX_ARCHIVE_COMPRESSION_RATIO for entry in entries):
            raise ValueError("Office 文档压缩比异常。")
        required_entry = "word/document.xml" if suffix == ".docx" else "xl/workbook.xml"
        if "[Content_Types].xml" not in names or required_entry not in names:
            raise ValueError(f"文件扩展名为 {suffix[1:].upper()}，但文档结构不匹配。")
    if suffix in {".txt", ".md"} and b"\x00" in content:
        raise ValueError("文本文件包含二进制内容。")


def parse_document(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if not content:
        raise ValueError("文件为空，请重新上传。")
    validate_document_signature(filename, content)

    if suffix == ".pdf":
        return _parse_pdf(content)
    if suffix == ".docx":
        return _parse_docx(content)
    if suffix == ".xlsx":
        return _parse_xlsx(content)
    if suffix in {".txt", ".md"}:
        text = content.decode("utf-8", errors="ignore").strip()
        if not text:
            raise ValueError("文本文件未解析出内容，请检查编码或文件内容。")
        return text
    raise ValueError("仅支持 PDF、DOCX、XLSX、TXT、MD 文件。")


def _parse_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError(f"PDF 页数不能超过 {MAX_PDF_PAGES} 页。")
        page_texts: dict[int, str] = {}
        pages_to_ocr: list[int] = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                page_texts[index] = text
            else:
                pages_to_ocr.append(index)
    except Exception as exc:
        raise ValueError(f"PDF 解析失败：{exc}") from exc

    if pages_to_ocr:
        page_texts.update(_ocr_pdf_pages(content, pages_to_ocr))

    parsed = "".join(
        f"\n\n--- 第 {index} 页 ---\n{text}"
        for index, text in sorted(page_texts.items())
        if text
    ).strip()
    if not parsed:
        raise ValueError("PDF 未解析出文本，OCR 也未识别到有效内容，请检查扫描清晰度。")
    return parsed


@lru_cache(maxsize=1)
def _get_ocr_engine() -> Any:
    try:
        from rapidocr import RapidOCR
    except ImportError as exc:
        raise ValueError("扫描版 PDF OCR 组件不可用，请重新安装应用后重试。") from exc
    return RapidOCR()


def _ocr_pdf_pages(content: bytes, page_numbers: list[int]) -> dict[int, str]:
    try:
        import fitz
    except ImportError as exc:
        raise ValueError("扫描版 PDF OCR 组件不可用，请重新安装应用后重试。") from exc

    try:
        document = fitz.open(stream=content, filetype="pdf")
        try:
            for page_number in page_numbers:
                _validate_ocr_page_size(document.load_page(page_number - 1))
            engine = _get_ocr_engine()
            page_texts = {}
            for page_number in page_numbers:
                page = document.load_page(page_number - 1)
                bitmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                result = engine(_pixmap_to_image(bitmap))
                text = _ocr_result_text(result)
                if text:
                    page_texts[page_number] = text
            return page_texts
        finally:
            close = getattr(document, "close", None)
            if close:
                close()
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"扫描版 PDF OCR 失败：{exc}") from exc


def _ocr_result_text(result: Any) -> str:
    texts = getattr(result, "txts", ())
    return "\n".join(text.strip() for text in texts if isinstance(text, str) and text.strip())


def _validate_ocr_page_size(page: Any) -> None:
    rect = getattr(page, "rect", None)
    width = getattr(rect, "width", None)
    height = getattr(rect, "height", None)
    if width is None or height is None:
        return
    pixels = int(float(width) * 2) * int(float(height) * 2)
    if pixels > MAX_OCR_PIXELS:
        raise ValueError(f"扫描 PDF 单页渲染像素不能超过 {MAX_OCR_PIXELS:,}。")


def _pixmap_to_image(bitmap: Any) -> Any:
    try:
        import numpy as np
    except ImportError as exc:
        raise ValueError("扫描版 PDF OCR 组件不可用，请重新安装应用后重试。") from exc
    return np.frombuffer(bitmap.samples, dtype=np.uint8).reshape(bitmap.height, bitmap.width, bitmap.n)


def _parse_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise ValueError(f"DOCX 解析失败：{exc}") from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    table_lines = []
    for table_index, table in enumerate(document.tables, start=1):
        table_lines.append(f"\n\n--- 表格 {table_index} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            table_lines.append(" | ".join(cells))

    parsed = "\n".join(paragraphs + table_lines).strip()
    if not parsed:
        raise ValueError("DOCX 未解析出文本，请检查文件内容。")
    return parsed


def _parse_xlsx(content: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ValueError("XLSX 解析组件不可用，请重新安装应用后重试。") from exc

    try:
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=False)
        sheet_lines = []
        row_count = 0
        for worksheet in workbook.worksheets:
            rows = []
            for row in worksheet.iter_rows(values_only=True):
                row_count += 1
                if row_count > MAX_XLSX_ROWS:
                    raise ValueError(f"XLSX 行数不能超过 {MAX_XLSX_ROWS:,} 行。")
                values = [str(value).strip() if value is not None else "" for value in row]
                if any(values):
                    rows.append(" | ".join(values))
            if rows:
                sheet_lines.append(f"--- 工作表：{worksheet.title} ---\n" + "\n".join(rows))
        workbook.close()
    except Exception as exc:
        raise ValueError(f"XLSX 解析失败：{exc}") from exc

    parsed = "\n\n".join(sheet_lines).strip()
    if not parsed:
        raise ValueError("XLSX 未解析出内容，请检查工作表数据。")
    return parsed
