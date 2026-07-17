import io
import asyncio
import json
import sys
import zipfile
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from openpyxl import Workbook

from backend.schemas import ApiConfig, WebSearchConfig
from backend.services.artifacts import build_output_files, extract_section, find_markdown_section, has_confirmed_content, infer_project_name, make_zip, markdown_line_diff, replace_markdown_section
from backend.services import document_parser
from backend.services.document_chunks import split_document_text
from backend.services.document_parser import parse_document
from backend.services.llm import create_agent
from backend.services.skill_loader import build_stage1_instructions, build_stage2_instructions, skill_source_label
from backend.services.web_search import WebSearchNotConfiguredError, build_search_context, tavily_search


def make_docx_bytes(text: str) -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def make_xlsx_bytes() -> bytes:
    buffer = io.BytesIO()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "评分表"
    worksheet.append(["评分项", "分值"])
    worksheet.append(["技术方案", 40])
    workbook.create_sheet("项目概况").append(["项目名称", "XLSX 项目"])
    workbook.save(buffer)
    return buffer.getvalue()


def test_parse_txt_document():
    assert parse_document("招标.txt", "项目名称：测试项目".encode("utf-8")) == "项目名称：测试项目"


def test_parse_docx_document():
    parsed = parse_document("招标.docx", make_docx_bytes("项目名称：DOCX 项目"))
    assert "DOCX 项目" in parsed


def test_parse_xlsx_document():
    parsed = parse_document("招标评分表.xlsx", make_xlsx_bytes())

    assert "工作表：评分表" in parsed
    assert "技术方案 | 40" in parsed
    assert "XLSX 项目" in parsed


def test_parse_legacy_doc_uses_converter(monkeypatch):
    monkeypatch.setattr(document_parser, "_parse_legacy_doc", lambda _: "项目名称：DOC 项目")

    assert parse_document("招标.doc", document_parser.DOC_SIGNATURE + b"legacy") == "项目名称：DOC 项目"


def test_bundled_legacy_doc_converter_is_preferred_in_frozen_agent(tmp_path, monkeypatch):
    converter_dir = tmp_path / "doc-converter"
    converter_dir.mkdir()
    converter = converter_dir / "program" / "soffice"
    converter.parent.mkdir()
    converter.write_text("", encoding="utf-8")
    (converter_dir / "converter.json").write_text(json.dumps({"executable": "program/soffice"}), encoding="utf-8")
    monkeypatch.setattr(document_parser, "_agent_resource_dir", lambda: tmp_path)
    monkeypatch.setattr(document_parser.sys, "frozen", True, raising=False)
    monkeypatch.delenv("AI_WORKBENCH_DOC_CONVERTER", raising=False)

    assert document_parser._find_legacy_doc_converter() == str(converter.resolve())


def test_legacy_doc_converter_ignores_runtime_executable_override(monkeypatch):
    monkeypatch.setenv("AI_WORKBENCH_DOC_CONVERTER", "/tmp/untrusted-converter")
    monkeypatch.setattr(document_parser, "_bundled_legacy_doc_converter", lambda: None)
    monkeypatch.setattr(document_parser.sys, "frozen", True, raising=False)

    assert document_parser._find_legacy_doc_converter() is None


def test_legacy_doc_converter_uses_restricted_environment(monkeypatch):
    captured = {}

    def fake_run(command, **kwargs):
        captured["command"] = command
        captured["options"] = kwargs
        return SimpleNamespace(stdout=b"DOC text", returncode=0)

    monkeypatch.setenv("HTTP_PROXY", "http://proxy.invalid")
    monkeypatch.setattr(document_parser, "_find_legacy_doc_converter", lambda: "/usr/bin/textutil")
    monkeypatch.setattr(document_parser.subprocess, "run", fake_run)

    assert document_parser._parse_legacy_doc(document_parser.DOC_SIGNATURE + b"legacy") == "DOC text"
    assert captured["options"]["start_new_session"] is True
    assert captured["options"]["cwd"] == captured["options"]["env"]["HOME"]
    assert "HTTP_PROXY" not in captured["options"]["env"]


def test_ocr_rejects_oversized_page_before_rendering(monkeypatch):
    class FakePage:
        rect = SimpleNamespace(width=10_000, height=10_000)

        def get_pixmap(self, matrix, alpha):
            raise AssertionError("oversized page must not render")

    class FakeDocument:
        def load_page(self, page_index):
            return FakePage()

    fake_fitz = SimpleNamespace(open=lambda stream, filetype: FakeDocument(), Matrix=lambda x, y: (x, y))
    monkeypatch.setitem(sys.modules, "fitz", fake_fitz)

    with pytest.raises(ValueError, match="渲染像素"):
        document_parser._ocr_pdf_pages(b"%PDF-1.7 fake", [1])


def test_split_document_text_preserves_content_at_section_boundaries():
    text = "前言\n\n--- 第 1 页 ---\n" + "内容" * 9 + "\n\n--- 第 2 页 ---\n尾部要求"

    chunks = split_document_text(text, max_characters=24)

    assert "".join(chunks) == text
    assert all(len(chunk) <= 24 for chunk in chunks)
    assert any("尾部要求" in chunk for chunk in chunks)


def test_replacing_markdown_section_keeps_other_sections_unchanged():
    content = "# 方案正文\n\n## 总体策略\n旧策略\n\n### 子项\n旧子项\n\n## 服务计划\n保持不变"

    section = find_markdown_section(content, "总体策略")
    revised = replace_markdown_section(content, section, "新策略\n\n### 子项\n新子项")

    assert section.body == "旧策略\n\n### 子项\n旧子项"
    assert "## 总体策略\n\n新策略" in revised
    assert "## 服务计划\n保持不变" in revised
    assert "旧策略" not in revised


def test_markdown_line_diff_marks_added_removed_and_unchanged_lines():
    lines = markdown_line_diff("# 标题\n旧内容\n保留内容", "# 标题\n新内容\n保留内容")

    assert lines == [
        {"kind": "unchanged", "content": "# 标题"},
        {"kind": "removed", "content": "旧内容"},
        {"kind": "added", "content": "新内容"},
        {"kind": "unchanged", "content": "保留内容"},
    ]


def test_xlsx_signature_error():
    with pytest.raises(ValueError, match="XLSX"):
        parse_document("招标.xlsx", b"not-a-zip")


def test_office_document_rejects_excessive_uncompressed_content(monkeypatch):
    monkeypatch.setattr(document_parser, "MAX_ARCHIVE_UNCOMPRESSED_BYTES", 16)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "x" * 20)
        archive.writestr("word/document.xml", "<document />")

    with pytest.raises(ValueError, match="解压后过大"):
        parse_document("招标.docx", buffer.getvalue())


def test_parse_scanned_pdf_uses_ocr(monkeypatch):
    class FakePage:
        def extract_text(self):
            return ""

    class FakeReader:
        pages = [FakePage()]

    monkeypatch.setattr(document_parser, "PdfReader", lambda _: FakeReader())
    monkeypatch.setattr(document_parser, "_ocr_pdf_pages", lambda _, pages: {pages[0]: "项目名称：OCR 招标文件"})

    parsed = parse_document("扫描招标.pdf", b"%PDF-1.7 fake")

    assert "第 1 页" in parsed
    assert "OCR 招标文件" in parsed


def test_ocr_pdf_pages_returns_rapidocr_text(monkeypatch):
    class FakePage:
        def get_pixmap(self, matrix, alpha):
            assert alpha is False
            return "bitmap"

    class FakeDocument:
        def load_page(self, page_index):
            assert page_index == 0
            return FakePage()

    fake_fitz = SimpleNamespace(open=lambda stream, filetype: FakeDocument(), Matrix=lambda x, y: (x, y))
    monkeypatch.setitem(sys.modules, "fitz", fake_fitz)
    monkeypatch.setattr(document_parser, "_get_ocr_engine", lambda: lambda _: SimpleNamespace(txts=("项目名称：OCR 项目", "")))
    monkeypatch.setattr(document_parser, "_pixmap_to_image", lambda _: "image")

    assert document_parser._ocr_pdf_pages(b"%PDF-1.7 fake", [1]) == {1: "项目名称：OCR 项目"}


def test_empty_file_error():
    with pytest.raises(ValueError, match="文件为空"):
        parse_document("empty.txt", b"")


def test_unsupported_file_error():
    with pytest.raises(ValueError, match="仅支持"):
        parse_document("presentation.pptx", b"abc")


def test_project_name_sanitizes_illegal_chars():
    assert infer_project_name("项目名称：A/B:C*D?E") == "ABCDE"


def test_artifacts_and_zip_generation():
    files = build_output_files(
        "项目名称：测试项目\n## 四、标书制作规范\n字体要求",
        "## 方案正文\n内容\n## 绘图提示词 + 专业图纸需求清单\n提示词",
    )
    assert "测试项目_设计方案.md" in files
    zipped = make_zip(files)
    with zipfile.ZipFile(io.BytesIO(zipped)) as archive:
        assert "测试项目_设计方案.md" in archive.namelist()


def test_extract_section_accepts_numbered_heading():
    section = extract_section("一、方案正文\n内容\n二、绘图提示词 + 专业图纸需求清单\n图纸内容", ["绘图提示词"])
    assert "图纸内容" in section


def test_extract_section_keeps_child_headings():
    section = extract_section(
        "## 方案正文\n内容\n## 绘图提示词 + 专业图纸需求清单\n说明\n### 一、绘图提示词\n提示词正文\n### 二、专业图纸需求清单\n清单正文\n## 标书制作规范\n规范",
        ["绘图提示词"],
    )
    assert "提示词正文" in section
    assert "清单正文" in section
    assert "标书制作规范" not in section


def test_artifacts_omit_untriggered_optional_files():
    files = build_output_files(
        "项目名称：测试项目\n项目概况：公共建筑\n设计范围：方案设计\n成果提交：总平面图、效果图、分析图",
        "## 方案正文\n功能分区、交通组织、景观节点设计。",
    )
    assert set(files) == {"测试项目_招标文件信息提取.md", "测试项目_设计方案.md"}


def test_unmentioned_specification_section_does_not_create_artifact():
    files = build_output_files(
        "项目名称：测试项目\n## 四、标书制作规范\n- 装订方式：未提及\n- 封面格式：未提及",
        "## 方案正文\n内容",
    )
    assert not has_confirmed_content("## 四、标书制作规范\n- 装订方式：未提及")
    assert "测试项目_标书制作规范.md" not in files


def test_openai_compatible_agent_accepts_base_url():
    agent = create_agent(
        ApiConfig(
            provider="DeepSeek",
            base_url="https://api.deepseek.com",
            api_key="test-key",
            model="deepseek-v4-pro",
        ),
        "test",
    )
    assert agent is not None
    assert agent.model.role_map["system"] == "system"


def write_test_skill(root):
    references = root / "references"
    references.mkdir()
    (root / "SKILL.md").write_text("# 外部测试 Skill\n", encoding="utf-8")
    (references / "可复用模块卡片.md").write_text("外部可复用模块卡片", encoding="utf-8")


def test_skill_loader_uses_bundled_skill_by_default(monkeypatch):
    monkeypatch.delenv("BID_DESIGN_WRITER_SKILL_DIR", raising=False)

    assert skill_source_label() == "bundled:bid_design_writer"
    instructions = build_stage1_instructions()
    assert "招标设计方案编写助手" in instructions


def test_skill_loader_accepts_explicit_external_override(tmp_path, monkeypatch):
    write_test_skill(tmp_path)
    monkeypatch.setenv("BID_DESIGN_WRITER_SKILL_DIR", str(tmp_path))

    assert skill_source_label() == str(tmp_path)
    assert "外部测试 Skill" in build_stage1_instructions()
    instructions = build_stage2_instructions()
    assert "外部可复用模块卡片" in instructions


def test_skill_loader_invalid_external_override_does_not_fallback(tmp_path, monkeypatch):
    missing = tmp_path / "missing-skill"
    monkeypatch.setenv("BID_DESIGN_WRITER_SKILL_DIR", str(missing))

    with pytest.raises(FileNotFoundError, match="外部覆盖路径"):
        build_stage1_instructions()


def test_skill_loader_uses_dynamic_stage2_instructions(monkeypatch):
    monkeypatch.delenv("BID_DESIGN_WRITER_SKILL_DIR", raising=False)

    instructions = build_stage2_instructions()
    assert "候选模块卡片" in instructions
    assert "12 章设计标模板参考" not in instructions


def test_bid_design_writer_regression_prompts_are_valid_json():
    fixture = Path(__file__).parent / "fixtures" / "bid_design_writer_test_prompts.json"
    prompts = json.loads(fixture.read_text(encoding="utf-8"))

    assert prompts
    assert all({"id", "prompt", "expected"} <= item.keys() for item in prompts)


def test_tavily_search_requires_api_key(monkeypatch):
    monkeypatch.setattr("backend.services.web_search.workbench_store.resolve_tavily_api_key", lambda user_id: None)
    monkeypatch.setattr("backend.services.web_search.workbench_store.get_web_search_config", lambda user_id: WebSearchConfig())

    with pytest.raises(WebSearchNotConfiguredError, match="TAVILY_API_KEY"):
        asyncio.run(tavily_search("test-user", "AI workbench"))


def test_tavily_search_calls_api_and_builds_context(monkeypatch):
    captured = {}

    class FakeResponse:
        def raise_for_status(self):
            return None

        def json(self):
            return {
                "results": [
                    {"title": "Result A", "url": "https://example.com/a", "content": "A summary"},
                    {"title": "Result B", "url": "https://example.com/b", "content": "B summary"},
                ]
            }

    class FakeClient:
        def __init__(self, timeout):
            assert timeout == 20

        async def __aenter__(self):
            return self

        async def __aexit__(self, exc_type, exc, tb):
            return False

        async def post(self, endpoint, headers, json):
            captured["endpoint"] = endpoint
            captured["headers"] = headers
            captured["json"] = json
            return FakeResponse()

    monkeypatch.setattr("backend.services.web_search.workbench_store.resolve_tavily_api_key", lambda user_id: "test-key")
    monkeypatch.setattr(
        "backend.services.web_search.workbench_store.get_web_search_config",
        lambda user_id: WebSearchConfig(has_key=True, source="db", max_results=2, search_depth="basic"),
    )
    monkeypatch.setattr("backend.services.web_search.httpx.AsyncClient", FakeClient)

    results = asyncio.run(tavily_search("test-user", "AI workbench"))
    context = build_search_context(results)

    assert captured["endpoint"] == "https://api.tavily.com/search"
    assert captured["headers"]["Authorization"] == "Bearer test-key"
    assert captured["json"]["query"] == "AI workbench"
    assert captured["json"]["max_results"] == 2
    assert results[0].title == "Result A"
    assert "[1] Result A" in context
    assert "https://example.com/b" in context
